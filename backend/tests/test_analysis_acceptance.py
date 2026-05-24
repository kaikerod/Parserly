from __future__ import annotations

import asyncio
import json
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from typing import Any
from uuid import uuid4

import httpx
import pytest
from fastapi import UploadFile
from starlette.datastructures import Headers

from app.core.config import Settings
from app.core.quotas import FREE_ANALYSIS_LIMIT
from app.schemas.analysis import AnalysisReport
from app.services import analysis_service as analysis_module
from app.services.analysis_service import (
    AIAnalysisResult,
    AIAnalysisUnavailable,
    AnalysisService,
    InvalidResumeFile,
    MAX_UPLOAD_BYTES,
    OPENROUTER_API_URL,
    QuotaExceeded,
    UserAnalysisReservation,
)


RESUME_TEXT = (
    "Joao Silva engenheiro de software Python FastAPI SQL AWS liderou projetos "
    "de APIs, automacao, monitoramento, testes, seguranca, dados, Docker, "
    "Kubernetes, CI CD e arquitetura cloud. Entregou reducao de custos, melhoria "
    "de desempenho, aumento de qualidade, colaboracao com produto, atendimento "
    "a usuarios, metricas operacionais e documentacao tecnica para equipes "
    "multifuncionais em ambientes de alta disponibilidade."
)


VALID_REPORT: dict[str, Any] = {
    "overall_score": 82,
    "categories": {
        "keywords": {"score": 84, "feedback": "Bom uso de termos tecnicos."},
        "formatting": {"score": 80, "feedback": "Formato legivel para ATS."},
        "structure": {"score": 81, "feedback": "Secoes bem organizadas."},
        "contact_info": {"score": 75, "feedback": "Contato suficiente."},
        "quantifiable_achievements": {
            "score": 87,
            "feedback": "Resultados mensuraveis aparecem no texto.",
        },
    },
    "recommendations": [
        {
            "priority": "high",
            "action": "Detalhar impacto por projeto.",
            "expected_impact": "Melhora a leitura de senioridade.",
        }
    ],
    "detected_role": "Engenheiro de Software",
}


def make_service(runtime_dir: Path) -> AnalysisService:
    return AnalysisService(
        db_session=object(),  # type: ignore[arg-type]
        settings=Settings(
            environment="test",
            upload_tmp_dir=str(runtime_dir),
            openrouter_api_key="test-openrouter-key",
            openrouter_model="primary-model",
            openrouter_fallback_model="fallback-model",
        ),
    )


def make_upload(
    filename: str,
    content: bytes,
    content_type: str,
    *,
    declared_size: int | None = None,
) -> UploadFile:
    return UploadFile(
        file=BytesIO(content),
        filename=filename,
        size=len(content) if declared_size is None else declared_size,
        headers=Headers({"content-type": content_type}),
    )


def build_docx_resume() -> bytes:
    pytest.importorskip("docx")
    from docx import Document

    document = Document()
    document.add_heading("Joao Silva", level=1)
    document.add_paragraph(RESUME_TEXT)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_resume() -> bytes:
    lines = [
        "Joao Silva - Engenheiro de Software",
        "Python FastAPI SQL AWS Docker Kubernetes CI CD observabilidade testes.",
        "Liderou APIs e automacao com reducao de custos e melhoria de desempenho.",
        "Colaborou com produto, dados, seguranca, usuarios, clientes e arquitetura cloud.",
        "Documentou metricas operacionais, qualidade, disponibilidade e resultados mensuraveis.",
    ]
    stream_lines = ["BT /F1 12 Tf 15 TL 72 720 Td"]
    for index, line in enumerate(lines):
        if index:
            stream_lines.append("T*")
        stream_lines.append(f"({escape_pdf_text(line)}) Tj")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1")

    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        (
            b"3 0 obj\n"
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\n"
            b"endobj\n"
        ),
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
        (
            b"5 0 obj\n<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream\nendobj\n"
        ),
    ]

    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for pdf_object in objects:
        offsets.append(len(output))
        output.extend(pdf_object)

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)


def escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def test_valid_docx_upload_extracts_readable_resume_text(runtime_dir: Path) -> None:
    service = make_service(runtime_dir)
    upload = make_upload(
        "resume.docx",
        build_docx_resume(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    temp_path = asyncio.run(service._write_upload_to_temp_file(upload))
    try:
        extracted = asyncio.run(service._extract_text(temp_path, "resume.docx"))
        quality = service._validate_extracted_text_quality(
            extracted.text,
            file_type=extracted.file_type,
            extraction_method=extracted.extraction_method,
            pages_processed=extracted.pages_processed,
        )
    finally:
        temp_path.unlink(missing_ok=True)

    assert extracted.file_type == ".docx"
    assert extracted.extraction_method == "python-docx"
    assert "FastAPI" in extracted.text
    assert quality.words >= 40


def test_valid_pdf_upload_extracts_readable_resume_text(runtime_dir: Path) -> None:
    pytest.importorskip("pdfplumber")
    service = make_service(runtime_dir)
    upload = make_upload("resume.pdf", build_pdf_resume(), "application/pdf")

    temp_path = asyncio.run(service._write_upload_to_temp_file(upload))
    try:
        extracted = asyncio.run(service._extract_text(temp_path, "resume.pdf"))
        quality = service._validate_extracted_text_quality(
            extracted.text,
            file_type=extracted.file_type,
            extraction_method=extracted.extraction_method,
            pages_processed=extracted.pages_processed,
        )
    finally:
        temp_path.unlink(missing_ok=True)

    assert extracted.file_type == ".pdf"
    assert extracted.extraction_method == "pdfplumber"
    assert extracted.pages_processed == 1
    assert quality.words >= 40


@pytest.mark.parametrize(
    ("upload", "expected_reason"),
    [
        (
            make_upload("resume.txt", b"plain text", "text/plain"),
            "unsupported_type",
        ),
        (
            make_upload("empty.pdf", b"", "application/pdf", declared_size=0),
            "insufficient_content",
        ),
        (
            make_upload(
                "large.pdf",
                b"%PDF-1.4",
                "application/pdf",
                declared_size=MAX_UPLOAD_BYTES + 1,
            ),
            "file_too_large",
        ),
        (
            make_upload("resume.pdf", b"%PDF-1.4", "text/plain"),
            "unsupported_type",
        ),
    ],
)
def test_upload_validation_rejects_invalid_files(
    runtime_dir: Path,
    upload: UploadFile,
    expected_reason: str,
) -> None:
    service = make_service(runtime_dir)

    with pytest.raises(InvalidResumeFile) as exc_info:
        asyncio.run(service._write_upload_to_temp_file(upload))

    assert exc_info.value.reason == expected_reason


def test_corrupted_pdf_is_rejected_before_ai_call(runtime_dir: Path) -> None:
    pytest.importorskip("pdfplumber")
    service = make_service(runtime_dir)
    corrupted_pdf = runtime_dir / "corrupted.pdf"
    corrupted_pdf.write_bytes(b"%PDF-1.4\nnot a complete pdf")

    with pytest.raises(InvalidResumeFile) as exc_info:
        asyncio.run(service._extract_text(corrupted_pdf, "corrupted.pdf"))

    assert exc_info.value.reason == "corrupted"


def test_raw_pdf_text_is_rejected_as_unreadable_payload(runtime_dir: Path) -> None:
    service = make_service(runtime_dir)
    raw_pdf_text = " ".join(
        [
            "1 0 obj stream endstream endobj xref trailer startxref",
            "/Type /Catalog /Pages /Font /Contents /Resources FlateDecode",
        ]
        * 6
    )

    with pytest.raises(InvalidResumeFile) as exc_info:
        service._validate_extracted_text_quality(
            raw_pdf_text,
            file_type=".pdf",
            extraction_method="pdfplumber",
            pages_processed=1,
        )

    assert exc_info.value.reason == "raw_pdf_content"


def test_openrouter_invalid_json_retries_with_fallback_model(
    runtime_dir: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    service = make_service(runtime_dir)
    requested_payloads: list[dict[str, Any]] = []
    responses = [
        openrouter_response("not valid json"),
        openrouter_response(json.dumps(VALID_REPORT)),
    ]

    class FakeAsyncClient:
        def __init__(self, *args: Any, **kwargs: Any) -> None:
            pass

        async def __aenter__(self) -> "FakeAsyncClient":
            return self

        async def __aexit__(self, *args: Any) -> None:
            return None

        async def post(self, url: str, **kwargs: Any) -> httpx.Response:
            assert url == OPENROUTER_API_URL
            requested_payloads.append(kwargs["json"])
            return responses.pop(0)

    async def no_sleep(seconds: float) -> None:
        return None

    monkeypatch.setattr(analysis_module.httpx, "AsyncClient", FakeAsyncClient)
    monkeypatch.setattr(analysis_module.asyncio, "sleep", no_sleep)

    result = asyncio.run(service._run_openrouter_attempts(RESUME_TEXT))

    assert result.model_used == "fallback-model"
    assert result.report.overall_score == 82
    assert [payload["model"] for payload in requested_payloads] == [
        "primary-model",
        "fallback-model",
    ]
    assert "ATENCAO" in requested_payloads[1]["messages"][0]["content"]


def test_openrouter_timeout_exhausts_attempt_budget(
    runtime_dir: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    service = make_service(runtime_dir)

    class TimeoutAsyncClient:
        def __init__(self, *args: Any, **kwargs: Any) -> None:
            pass

        async def __aenter__(self) -> "TimeoutAsyncClient":
            return self

        async def __aexit__(self, *args: Any) -> None:
            return None

        async def post(self, url: str, **kwargs: Any) -> httpx.Response:
            raise httpx.ReadTimeout("OpenRouter timed out")

    async def no_sleep(seconds: float) -> None:
        return None

    monkeypatch.setattr(analysis_module.httpx, "AsyncClient", TimeoutAsyncClient)
    monkeypatch.setattr(analysis_module.asyncio, "sleep", no_sleep)

    with pytest.raises(AIAnalysisUnavailable):
        asyncio.run(service._run_openrouter_attempts(RESUME_TEXT))


class PersistResult:
    def __init__(self, row: tuple[int, ...]) -> None:
        self.row = row

    def one_or_none(self) -> tuple[int, ...]:
        return self.row


class PersistDbSession:
    def __init__(self, row: tuple[int, ...]) -> None:
        self.row = row
        self.statements: list[object] = []
        self.committed = False
        self.rolled_back = False

    def add(self, instance: object) -> None:
        self.added = instance

    async def execute(self, statement: object) -> PersistResult:
        self.statements.append(statement)
        return PersistResult(self.row)

    async def commit(self) -> None:
        self.committed = True

    async def rollback(self) -> None:
        self.rolled_back = True

    async def refresh(self, instance: object) -> None:
        instance.id = uuid4()
        instance.created_at = datetime(2026, 5, 7, 12, 0, tzinfo=UTC)


class ReserveDbSession:
    def __init__(self, row: tuple[int, ...]) -> None:
        self.row = row
        self.statements: list[object] = []
        self.committed = False
        self.rolled_back = False

    async def refresh(self, instance: object) -> None:
        return None

    async def execute(self, statement: object) -> PersistResult:
        self.statements.append(statement)
        return PersistResult(self.row)  # type: ignore[arg-type]

    async def commit(self) -> None:
        self.committed = True

    async def rollback(self) -> None:
        self.rolled_back = True


class UserWithExpiredCreditsAfterCommit:
    def __init__(self, db_session: ReserveDbSession) -> None:
        self.id = uuid4()
        self.analyses_used = FREE_ANALYSIS_LIMIT
        self._paid_analysis_credits = 10
        self._db_session = db_session

    @property
    def paid_analysis_credits(self) -> int:
        if self._db_session.committed:
            raise AssertionError("paid_analysis_credits was read after commit")
        return self._paid_analysis_credits

    @paid_analysis_credits.setter
    def paid_analysis_credits(self, value: int) -> None:
        self._paid_analysis_credits = value


def test_reserve_user_analysis_consumes_paid_credit_after_free_quota(
    runtime_dir: Path,
) -> None:
    db_session = ReserveDbSession((FREE_ANALYSIS_LIMIT, 9))
    service = AnalysisService(
        db_session=db_session,  # type: ignore[arg-type]
        settings=Settings(environment="test", upload_tmp_dir=str(runtime_dir)),
    )
    user = UserWithExpiredCreditsAfterCommit(db_session)

    reservation = asyncio.run(service._reserve_user_analysis(user))  # type: ignore[arg-type]

    assert reservation == UserAnalysisReservation(
        analyses_used=FREE_ANALYSIS_LIMIT,
        consumed_paid_credit=True,
    )
    assert user._paid_analysis_credits == 10
    assert db_session.committed is False
    assert db_session.rolled_back is False
    assert db_session.statements == []


def test_master_admin_user_analysis_does_not_consume_quota_or_paid_credits(
    runtime_dir: Path,
) -> None:
    db_session = ReserveDbSession((FREE_ANALYSIS_LIMIT, 0))
    service = AnalysisService(
        db_session=db_session,  # type: ignore[arg-type]
        settings=Settings(environment="test", upload_tmp_dir=str(runtime_dir)),
    )
    user = SimpleNamespace(
        id=uuid4(),
        email="kaikevinicius789@gmail.com",
        analyses_used=FREE_ANALYSIS_LIMIT,
        paid_analysis_credits=0,
    )

    reservation = asyncio.run(service._reserve_user_analysis(user))  # type: ignore[arg-type]
    analyses_used = asyncio.run(
        service._commit_user_analysis_usage(user, reservation)  # type: ignore[arg-type]
    )

    assert reservation == UserAnalysisReservation(
        analyses_used=FREE_ANALYSIS_LIMIT,
        consumed_paid_credit=False,
        unlimited=True,
    )
    assert analyses_used == FREE_ANALYSIS_LIMIT
    assert user.analyses_used == FREE_ANALYSIS_LIMIT
    assert user.paid_analysis_credits == 0
    assert db_session.statements == []


def test_standard_user_without_remaining_analyses_still_hits_quota_limit(
    runtime_dir: Path,
) -> None:
    db_session = ReserveDbSession((FREE_ANALYSIS_LIMIT, 0))
    service = AnalysisService(
        db_session=db_session,  # type: ignore[arg-type]
        settings=Settings(environment="test", upload_tmp_dir=str(runtime_dir)),
    )
    user = SimpleNamespace(
        id=uuid4(),
        email="person@example.com",
        analyses_used=FREE_ANALYSIS_LIMIT,
        paid_analysis_credits=0,
    )

    with pytest.raises(QuotaExceeded):
        asyncio.run(service._reserve_user_analysis(user))  # type: ignore[arg-type]


def test_persist_analysis_uses_reserved_paid_credit(runtime_dir: Path) -> None:
    db_session = PersistDbSession((FREE_ANALYSIS_LIMIT, 9))
    service = AnalysisService(
        db_session=db_session,  # type: ignore[arg-type]
        settings=Settings(environment="test", upload_tmp_dir=str(runtime_dir)),
    )
    user = SimpleNamespace(
        id=uuid4(),
        analyses_used=FREE_ANALYSIS_LIMIT,
        paid_analysis_credits=10,
    )
    ai_result = AIAnalysisResult(
        report=AnalysisReport.model_validate(VALID_REPORT),
        model_used="test-model",
    )

    result = asyncio.run(
        service._persist_analysis(
            user=user,  # type: ignore[arg-type]
            filename="resume.pdf",
            ai_result=ai_result,
            guest_analyses_used=None,
            user_reservation=UserAnalysisReservation(
                analyses_used=FREE_ANALYSIS_LIMIT,
                consumed_paid_credit=True,
            ),
        )
    )

    assert result.analyses_used == FREE_ANALYSIS_LIMIT
    assert user.paid_analysis_credits == 9
    assert db_session.committed is True
    assert db_session.rolled_back is False
    assert len(db_session.statements) == 1
    assert "paid_analysis_credits" in str(db_session.statements[0])


def openrouter_response(content: str) -> httpx.Response:
    return httpx.Response(
        200,
        json={"choices": [{"message": {"content": content}}]},
        request=httpx.Request("POST", OPENROUTER_API_URL),
    )
